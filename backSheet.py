from docx import Document
from docx.shared import Pt,Inches,Cm
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT,WD_TABLE_ALIGNMENT
import pandas as pd

df=pd.read_excel(f"Student Details.xlsx")

doc=Document()
section=doc.sections[-1]
section.page_height=Cm(29.7)
section.page_width=Cm(21)
section.left_margin=Cm(1.0)
section.right_margin=Cm(1.0)
section.top_margin=Cm(0.8)
section.bottom_margin=Cm(0.4)

new_width,new_height=section.page_height,section.page_width
section.orientation=WD_ORIENTATION.LANDSCAPE
section.page_width=new_width
section.page_height=new_height
sectPr=section._sectPr

y='2021-2022'                                                      # academic year

try:
    k=0
    while(True):
        Name=df.iloc[k].Name
        table1=doc.add_table(rows=2,cols=1,style='Table Grid')
        table1.rows[1].height=Cm(17)
        row1=table1.rows[0]
        table1.cell(0,0).paragraphs[0].add_run("\n").font.size=Pt(7)
        heading=row1.cells[0].paragraphs[0].add_run("Millennium Model School Mandi Bamora")
        heading.bold=True
        heading.font.size=Pt(18)
        heading1=row1.cells[0].paragraphs[0].add_run(f"\nReport Card ({y})\n")
        heading1.bold=True
        heading1.italic=True
        heading1.font.size=Pt(15)
        row2=table1.rows[1]
        table1.cell(0,0).paragraphs[0].paragraph_format.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER

        dob=pd.to_datetime(df.iloc[k].DOB).date().strftime('%d-%m-%Y')

        r2c1=table1.cell(1,0)
        detailtable=r2c1.add_table(rows=1,cols=2)
        detailcell1=detailtable.cell(0,0)
        detailcell2=detailtable.cell(0,1)
        detail1=detailcell1.paragraphs[0].add_run(f"Student's Name :  {Name}\nFather's Name   :  Mr. {df.iloc[k].Father}")
        detail2=detailcell2.paragraphs[0].add_run(f"\t\t\t\t\t           Class     :  {df.iloc[k].Class}\n\t\t\t\t\tDate of Birth :  {dob}")
        detail1.font.size=detail2.font.size=Pt(13)
        table2=r2c1.add_table(rows=1,cols=2)
        leftcell=table2.cell(0,0)
        rightcell=table2.cell(0,1)
        table3=leftcell.add_table(rows=7,cols=4)
        for row in table3.rows:
            row.height=Cm(0.9)
        table3.rows[0].height=Cm(1.85)
        t300=table3.cell(0,0)
        l1=['Valuation of Educational Field (Grade)','English','Hindi','Mathematics','Science','Social Science','Sanskrit']
        for i in range(7):
            table3.cell(i,0).paragraphs[0].add_run(l1[i]).bold=True
            for j in range(4):
                table3.cell(i,j).paragraphs[0].paragraph_format.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
                table3.cell(i,j).vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        table3.cell(0,1).paragraphs[0].add_run('Half Yearly Exam').bold=True
        table3.cell(0,2).paragraphs[0].add_run('Yearly Exam').bold=True
        table3.cell(0,3).paragraphs[0].add_run('Annual Result').bold=True
        table4=rightcell.add_table(rows=7,cols=4)
        for row in table4.rows:
            row.height=Cm(0.9)
        table4.rows[0].height=Cm(1.85)
        for cell in table4.columns[0].cells:
            cell.width=Cm(4)
        l2=['Valuation of Academic Field (Grade)','Library','Cultural Activities','Sharpness','Creativity','Yoga','Sports']
        for i in range(7):
            table4.cell(i,0).paragraphs[0].add_run(l2[i]).bold=True
            for j in range(4):
                table4.cell(i,j).paragraphs[0].paragraph_format.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
                table4.cell(i,j).vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        l3=['Valuation of Social and Personal Values Field (Grade)','Discipline','Punctuality','Cleanliness','Honesty','Truthfulness','Attitude']
        for i in range(7):
            table4.cell(i,2).paragraphs[0].add_run(l3[i]).bold=True
        table4.cell(0,1).paragraphs[0].add_run('Annual Result').bold=True
        table4.cell(0,3).paragraphs[0].add_run('Annual Result').bold=True
        table3.style='Table Grid'
        table4.style='Table Grid'
        table5=table1.cell(1,0).add_table(rows=1,cols=1)
        table5.style='Table Grid'
        table5.rows[0].height=Cm(0.6)
        table5.rows[0].cells[0].width=Cm(26.84)
        table5.alignment=WD_TABLE_ALIGNMENT.CENTER
        table5.cell(0,0).paragraphs[0].add_run(f"Total Days of Teaching : 270\t\t\t\t\t\t\t        Student's Attendence : {int(df.iloc[k].Attendence)}")
        table5.cell(0,0).vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        table6=table1.cell(1,0).add_table(rows=1,cols=4)
        gs=table6.cell(0,0).add_paragraph('\nGrade System')
        gs.paragraph_format.space_after=Pt(0)
        gs.paragraph_format.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
        table7=table6.cell(0,0).add_table(rows=5,cols=3)
        table7.style='Table Grid'
        l4=[['75-100%','60-75%','45-60%','33-45%','0-30%'],['A','B','C','D','E'],['Excellent','Very Good','Good','To be taken care','Poor']]
        for i in range(5):
            for j in range(3):
                table7.cell(i,j).paragraphs[0].add_run(l4[j][i])
                table7.cell(i,j).paragraphs[0].paragraph_format.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
                table7.cell(i,j).vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for cell in table7.columns[1].cells:
            cell.width=Cm(1)
        for cell in table7.columns[2].cells:
            cell.width=Cm(3.5)
        table6.cell(0,1).width=Cm(11.5)
        table6.cell(0,2).width=Cm(3.5)
        table6.cell(0,3).width=Cm(3.5)
        table6.rows[0].height=Cm(4)
        t=table6.cell(0,1).add_paragraph("\n")
        t.paragraph_format.line_spacing=Pt(18)

        annuualgrade=df.iloc[k].Grade
        if(annuualgrade=='A'):
            remarks=l4[2][0]
        elif(annuualgrade=='B'):
            remarks=l4[2][1]
        elif(annuualgrade=='C'):
            remarks=l4[2][2]
        elif(annuualgrade=='D'):
            remarks=l4[2][3]
        elif(annuualgrade=='E'):
            remarks=l4[2][4]
        else:
            c=1/0

        table6.cell(0,1).paragraphs[1].add_run(f"\n\tTeacher's Remark         :  {remarks}\n\tAnnual Result (Grade) :  {annuualgrade}\n\tThe student has been promoted in class {str(int(df.iloc[k].Class[0])+1)+'th'}").font.size=Pt(13)
        ct=table6.cell(0,2).paragraphs[0].add_run('\n\n\n\n\n\n   Class Teacher')
        ct.bold=True
        ct.font.size=Pt(12.5)
        hm=table6.cell(0,3).paragraphs[0].add_run('\n\n\n\n\n\n\tHeadmaster')
        hm.bold=True
        hm.font.size=Pt(12.5)
        hm.line_spacing=Pt(13)


        halfyearly=df.iloc[k].Half.split()
        for i in range(1,7):
            table3.cell(i,1).paragraphs[0].add_run(f'{halfyearly[i-1]}')

        yearly=df.iloc[k].Yearly.split()
        for i in range(1,7):
            table3.cell(i,2).paragraphs[0].add_run(f'{yearly[i-1]}')

        for i in range(1,7):
            table3.cell(i,3).paragraphs[0].add_run(f'{max(halfyearly[i-1],yearly[i-1])}')

        academic=df.iloc[k].Academic.split()
        for i in range(1,7):
            table4.cell(i,1).paragraphs[0].add_run(f'{academic[i-1]}')

        social=df.iloc[k].Social.split()
        for i in range(1,7):
            table4.cell(i,3).paragraphs[0].add_run(f'{social[i-1]}')
        
        k=k+1
        df.iloc[k].Name
        doc.add_page_break()
except ZeroDivisionError:
    print(f'\n-------------- Error : Check Annual Grade of student no. {k+1} ----------------\n')
    exit()

except IndexError:
    pass
except:
    print("There is some error. Please check your document carefully.")

doc.save(f'Marksheets (back page) {y}.docx')
print(f'\nYour file has been saved with file name - Marksheets (back page) {y}\n')