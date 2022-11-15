from os import truncate
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT,WD_TABLE_ALIGNMENT
from docx.shared import Pt,Inches,Cm
from num2words import num2words
import pandas as pd
df=pd.read_excel(f"Student Details.xlsx")
y='2021-2022'                                                     # academic year
def datetoWord(d):
    dic={1:'January',2:'February',3:'March',4:'April',5:'May',6:'June',7:'July',8:'August',9:'September',10:'October',11:'November',12:'December'}
    w=num2words(int(d[6:])).title().split()
    w1=w[:2]+w[3:]
    fw=''
    for i in w1:
        fw=fw+i+' '
    yy=num2words(int(d[0:2])).title() +'  ' + dic[int(d[3:5])] +'\n                    (in words)\t\t'+ fw.title()
    return yy
doc=Document()
section=doc.sections[-1]
section.page_height=Cm(29.7)
section.page_width=Cm(21)
section.left_margin=Cm(1.0)
section.right_margin=Cm(1.0)
section.top_margin=Cm(0.8)
section.bottom_margin=Cm(0.4)

new_width,new_height=section.page_height,section.page_width
section.orientation=WD_ORIENTATION.LANDSCAPE
section.page_width=new_width
section.page_height=new_height
sectPr=section._sectPr
cols=sectPr.xpath('./w:cols')[0]
cols.set(qn('w:num'),'2')

#year=input("Enter the year : ")
#s=int(input("Enter the number of students : "))
try:
    k=0
    while(True):
        dob=pd.to_datetime(df.iloc[k].DOB).date().strftime('%d-%m-%Y')
        table1=doc.add_table(rows=1,cols=1,style='Table Grid')
        table1.rows[0].height=Cm(19.3)
        c=table1.cell(0,0).paragraphs[0].add_run('''\n
    If you want to build a nation, educate the children. If you\n
    want to make the nation strong and integrated, educate\n
    the children. Any investment, any expenditure made on\n
    the education of the children is  the  investment  on  the\n
    nation. The children and the nation are safer and secure\n
    in the hands of right sort of teachers.\n\n''')
        c1=table1.cell(0,0).paragraphs[0].add_run('\n\n  Education is the only panacea for all kinds of our evils.\n\n\n\n\n\n\n\n\n')
        #doc.add_paragraph('\n')
        c.font.size=Pt(14.5)
        c1.font.size=Pt(14.5)
        c1.bold=True

        table2=doc.add_table(rows=1,cols=1,style='Table Grid')
        table2.rows[0].height=Cm(19.3)

        table2.cell(0,0).paragraphs[0].add_run('\n\n').font.size=Pt(15)
        #m=table2.cell(0,0).paragraphs[-1].add_run('Millennium Model School\n')
        #m.font.size=Pt(20)
        #m.bold=True
            
        run=table2.cell(0,0).paragraphs[0].add_run()
        p=run.add_picture('logo.jpg',width=Cm(9.5),height=Cm(5.7))
        table2.cell(0,0).paragraphs[0].alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
        #pc=table2.cell(0,0).paragraphs[0].add_run(f'\n\nREPORT CARD {year}')
        #pc.bold=True
        #pc.font.size=Pt(17)

        table2.cell(0,0).add_paragraph('\n\n')
        det=table2.cell(0,0).paragraphs[1].add_run(f'''\n
                    Student's Name        :  \t{df.iloc[k].Name}\n
                    Class                            :  \t{df.iloc[k].Class}\n
                    Scholar No.                :  \t{df.iloc[k].ScholarNo}\n
                    Father's Name          :  \tMr. {df.iloc[k].Father}\n
                    Mother's Name        :  \tMrs. {df.iloc[k].Mother}\n
                    Date of Birth             :  \t{dob}\n
                    Date of birth             :  \t{datetoWord(dob)}\n 
                    SSMID No.                 :  \t{df.iloc[k].SSMID}\n
                    Aadhar No.               :  \t{df.iloc[k].Aadhar}\n''')
        det.font.size=Pt(13)
        k=k+1
        df.iloc[k].Name        #it will not add page break at last
        doc.add_page_break()
except IndexError:
    pass
except:
    print("There is some error. Please check your document carefully.")

doc.save(f'Marksheets (front page) {y}.docx')
print(f'\nYour file has been saved with file name - Marksheets (front page) {y}\n')